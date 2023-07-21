from tkinter import *
from tkinter import messagebox
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy import Column, Integer, String, Float, ForeignKey
from sqlalchemy.orm import relationship
from tkinter import ttk
from docx import Document

# Crear la base de datos
engine = create_engine('mysql+mysqlconnector://root:123456@localhost/learning_tracker_h2')
Base = declarative_base()

# Definir las clases de la base de datos
class Estudiante(Base):
    __tablename__ = 'estudiantes'

    id = Column(Integer, primary_key=True)
    nombre = Column(String)
    apellido = Column(String)
    edad = Column(Integer)
    codigo_estudiante = Column(String, unique=True)
    promedio = Column(Float)

    calificaciones = relationship("Calificacion", back_populates="estudiante")

class Materia(Base):
    __tablename__ = 'materias'

    id = Column(Integer, primary_key=True)
    nombre = Column(String, unique=True)

    calificaciones = relationship("Calificacion", back_populates="materia")

class Calificacion(Base):
    __tablename__ = 'calificaciones'

    id = Column(Integer, primary_key=True)
    id_estudiante = Column(Integer, ForeignKey('estudiantes.id'))
    id_materia = Column(Integer, ForeignKey('materias.id'))
    calificacion = Column(Float)

    estudiante = relationship("Estudiante", back_populates="calificaciones")
    materia = relationship("Materia", back_populates="calificaciones")

# Crear las tablas en la base de datos
Base.metadata.create_all(engine)

def agregar_estudiante():
    # Obtener los valores ingresados en el formulario
    nombre = entry_nombre.get()
    apellido = entry_apellido.get()
    edad = entry_edad.get()
    codigo_estudiante = entry_codigo_estudiante.get()

    # Validar que se hayan ingresado todos los datos
    if not nombre or not apellido or not edad or not codigo_estudiante:
        messagebox.showerror("Error", "Por favor, complete todos los campos.")
        return

    # Crear el motor de la base de datos y la sesión
    engine = create_engine('mysql+mysqlconnector://root:123456@localhost/learning_tracker_h2')
    Session = sessionmaker(bind=engine)
    session = Session()

    # Verificar si el estudiante ya existe en la base de datos
    estudiante_existente = session.query(Estudiante).filter(Estudiante.codigo_estudiante == codigo_estudiante).first()
    if estudiante_existente:
        messagebox.showerror("Error", "El estudiante ya existe en la base de datos.")
        session.close()
        return

    # Crear una instancia de Estudiante con los datos ingresados
    estudiante = Estudiante(nombre=nombre, apellido=apellido, edad=edad, codigo_estudiante=codigo_estudiante, promedio=0)

    # Agregar el estudiante a la base de datos
    session.add(estudiante)

    # Agregar las calificaciones iniciales con valor cero
    materias = session.query(Materia).all()
    for materia in materias:
        calificacion_nueva = Calificacion(calificacion=0)
        estudiante.calificaciones.append(calificacion_nueva)
        materia.calificaciones.append(calificacion_nueva)
        session.add(calificacion_nueva)

    # Commit para guardar todos los cambios
    session.commit()

    # Cerrar la sesión
    session.close()

    # Limpiar los campos del formulario
    entry_nombre.delete(0, END)
    entry_apellido.delete(0, END)
    entry_edad.delete(0, END)
    entry_codigo_estudiante.delete(0, END)

    # Mostrar mensaje de éxito
    messagebox.showinfo("Éxito", "El estudiante se ha agregado correctamente.")

def mostrar_estudiantes():
    # Crear el motor de la base de datos y la sesión
    engine = create_engine('mysql+mysqlconnector://root:123456@localhost/learning_tracker_h2')
    Session = sessionmaker(bind=engine)
    session = Session()

    # Obtener todos los estudiantes de la base de datos
    estudiantes = session.query(Estudiante).all()

    # Limpiar la lista de estudiantes
    tree.delete(*tree.get_children())

    # Insertar los estudiantes en la lista
    for estudiante in estudiantes:
        tree.insert("", "end", values=(estudiante.id, estudiante.nombre, estudiante.apellido, estudiante.edad, estudiante.codigo_estudiante, estudiante.promedio))

    # Cerrar la sesión
    session.close()

def generar_informe():
    # Obtener los estudiantes seleccionados en la lista
    seleccion = tree.selection()
    if not seleccion:
        messagebox.showerror("Error", "Por favor, seleccione al menos un estudiante.")
        return

    # Crear el informe en un documento de Word
    document = Document()
    document.add_heading('Informe de Rendimiento Académico', level=1)

    # Crear el motor de la base de datos y la sesión
    engine = create_engine('mysql+mysqlconnector://root:123456@localhost/learning_tracker_h2')
    Session = sessionmaker(bind=engine)
    session = Session()

    # Generar el informe para cada estudiante seleccionado
    for item in seleccion:
        estudiante_id = tree.item(item, "values")[0]
        estudiante = session.query(Estudiante).get(estudiante_id)

        document.add_heading(f'Estudiante: {estudiante.nombre} {estudiante.apellido}', level=2)

        # Obtener las calificaciones del estudiante
        calificaciones = session.query(Calificacion).join(Materia).filter(Calificacion.id_estudiante == estudiante.id).all()

        if calificaciones:
            table = document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            table.autofit = False

            # Encabezados de la tabla
            table.cell(0, 0).text = 'Materia'
            table.cell(0, 1).text = 'Calificación'
            table.cell(0, 2).text = 'Promedio'

            # Contador para el promedio
            contador = 0

            # Datos de calificaciones en la tabla
            for i, calificacion in enumerate(calificaciones):
                materia = session.query(Materia).get(calificacion.id_materia)

                table.add_row()
                table.cell(i+1, 0).text = materia.nombre
                table.cell(i+1, 1).text = str(calificacion.calificacion)

                contador += calificacion.calificacion

            promedio = contador / len(calificaciones)

            # Celda para mostrar el promedio
            table.add_row()
            table.cell(len(calificaciones)+1, 2).text = f'Promedio: {promedio}'

        else:
            document.add_paragraph("No hay calificaciones para mostrar.")

        document.add_paragraph("")  # Espacio entre estudiantes

    # Guardar el documento en formato Word
    document.save("informe.docx")

    # Cerrar la sesión
    session.close()

  
# Crear la ventana principal
ventana = Tk()
ventana.title("Learning Tracker")
ventana.geometry("800x600")

# Crear la tabla para mostrar los estudiantes
tree = ttk.Treeview(ventana)
tree["columns"] = ("ID", "Nombre", "Apellido", "Edad", "Código", "Promedio")

# Configurar las columnas de la tabla
tree.column("#0", width=0, stretch=NO)
tree.column("ID", anchor=CENTER, width=40)
tree.column("Nombre", anchor=W, width=120)
tree.column("Apellido", anchor=W, width=120)
tree.column("Edad", anchor=CENTER, width=40)
tree.column("Código", anchor=CENTER, width=100)
tree.column("Promedio", anchor=CENTER, width=60)

# Encabezado de las columnas
tree.heading("#0", text="", anchor=CENTER)
tree.heading("ID", text="ID", anchor=CENTER)
tree.heading("Nombre", text="Nombre", anchor=W)
tree.heading("Apellido", text="Apellido", anchor=W)
tree.heading("Edad", text="Edad", anchor=CENTER)
tree.heading("Código", text="Código", anchor=CENTER)
tree.heading("Promedio", text="Promedio", anchor=CENTER)

# Barra de desplazamiento vertical
scrollbar = ttk.Scrollbar(ventana, orient="vertical", command=tree.yview)
scrollbar.pack(side=RIGHT, fill=Y)
tree.configure(yscrollcommand=scrollbar.set)

# Agregar la tabla a la ventana
tree.pack(pady=10)

# Botones
frame_botones = Frame(ventana)
frame_botones.pack()

btn_agregar = Button(frame_botones, text="Agregar Estudiante", command=agregar_estudiante)
btn_agregar.pack(side=LEFT, padx=10)

btn_mostrar = Button(frame_botones, text="Mostrar Estudiantes", command=mostrar_estudiantes)
btn_mostrar.pack(side=LEFT, padx=10)

btn_informe = Button(frame_botones, text="Generar Informe", command=generar_informe)
btn_informe.pack(side=LEFT, padx=10)

# Formulario para agregar estudiantes
frame_formulario = Frame(ventana)
frame_formulario.pack()

label_nombre = Label(frame_formulario, text="Nombre:")
label_nombre.grid(row=0, column=0, padx=5, pady=5, sticky=E)

entry_nombre = Entry(frame_formulario)
entry_nombre.grid(row=0, column=1, padx=5, pady=5)

label_apellido = Label(frame_formulario, text="Apellido:")
label_apellido.grid(row=1, column=0, padx=5, pady=5, sticky=E)

entry_apellido = Entry(frame_formulario)
entry_apellido.grid(row=1, column=1, padx=5, pady=5)

label_edad = Label(frame_formulario, text="Edad:")
label_edad.grid(row=2, column=0, padx=5, pady=5, sticky=E)

entry_edad = Entry(frame_formulario)
entry_edad.grid(row=2, column=1, padx=5, pady=5)

label_codigo_estudiante = Label(frame_formulario, text="Código:")
label_codigo_estudiante.grid(row=3, column=0, padx=5, pady=5, sticky=E)

entry_codigo_estudiante = Entry(frame_formulario)
entry_codigo_estudiante.grid(row=3, column=1, padx=5, pady=5)

# Ejecutar la ventana principal
ventana.mainloop()
